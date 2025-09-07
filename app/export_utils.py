# app/export_utils.py
import io
import os
from uuid import uuid4
from werkzeug.utils import secure_filename
from flask import current_app
from docx import Document
from docx.shared import Pt
from datetime import datetime
from openpyxl import Workbook
from openpyxl.utils import get_column_letter

def make_filename(prefix, ext):
    name = f"{prefix}-{uuid4().hex}.{ext}"
    return secure_filename(name)

def save_bytes_to_uploads(content_bytes: bytes, filename: str):
    """
    Сохраняет байты в папку uploads и возвращает путь (имя файла на диске) и размер.
    """
    upload_folder = current_app.config["UPLOAD_FOLDER"]
    os.makedirs(upload_folder, exist_ok=True)
    path = os.path.join(upload_folder, filename)
    with open(path, "wb") as f:
        f.write(content_bytes)
    size = os.path.getsize(path)
    return filename, size

def generate_certificate_docx(user, course, issuer_name="Московский университет"):
    """
    Возвращает bytes (docx) для сертификата ученика user по course
    """
    doc = Document()

    # простой стиль заголовка
    h = doc.add_heading(level=1)
    run = h.add_run("Сертификат о прохождении курса")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph("")  # пустая

    # студент и курс
    p = doc.add_paragraph()
    run = p.add_run(f"Настоящий сертификат подтверждает, что ")
    run.font.size = Pt(12)
    run_b = p.add_run(f"{user.username} ({user.email})")
    run_b.bold = True
    run_b.font.size = Pt(12)
    p.add_run(f" успешно завершил(а) курс ").font.size = Pt(12)
    p.add_run(f"{course.title}").bold = True

    doc.add_paragraph("")  # пустая

    p2 = doc.add_paragraph()
    p2.add_run(f"Дата выдачи: {datetime.utcnow().strftime('%d.%m.%Y')}")
    p2.add_run("\n\n")
    p2.add_run(f"Выдано: {issuer_name}")

    doc.add_paragraph("")
    doc.add_paragraph("Подпись: ____________________")

    # Сохраняем в байты
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def generate_progress_xlsx(course_id, rows):
    """
    rows: iterable of tuples (student_name, email, lesson_title, status, score, completed_at)
    Возвращает bytes xlsx
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Progress"

    headers = ["Student", "Email", "Lesson", "Status", "Score", "Completed At"]
    ws.append(headers)

    for r in rows:
        # completed_at -> string
        ca = r[5].strftime("%Y-%m-%d %H:%M") if r[5] else ""
        ws.append([r[0], r[1], r[2], r[3], float(r[4]) if r[4] is not None else "", ca])

    # auto width
    for i, col in enumerate(ws.columns, 1):
        max_length = 0
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[get_column_letter(i)].width = min(max_length + 2, 50)

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()